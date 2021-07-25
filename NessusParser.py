#/usr/bin/python3.8
import csv,sys,docx,copy
import os
from docx.enum.text import WD_COLOR_INDEX

# def getColor(elem): #restituisce un colore della libreria python-docx 
#     if elem == "Medium":
#         return WD_COLOR_INDEX.YELLOW
#     elif elem == "High":
#         return WD_COLOR_INDEX.RED
#     elif elem == "Critical":
#         return WD_COLOR_INDEX.DARK_RED

def getVulnerableIps(csv, str): #returns the ip of a row if it has an id equal to "str"
    result = set()
    for entry in csv:
        if entry[0] == str: #id comparison
            result.add(entry[4]) #insert the id (without duplicates)
            #every host can have duplicates because there can be the same ip in multiple lines
    return result

def writeOnWord(csvlist, ids, documentPar):
    head = documentPar.insert_paragraph_before("Analysing Subnet " + Singlefile) #TODO aggiunti font
    print("Analysing Subnet " + Singlefile)
    mediumVulns = []
    critVulns = []
    highVulns = []
    
    for elem in csvlist:
        if not elem[0] in ids: #if the id is not in the set
            ids.add(elem[0])#i'll add it to the set
            if elem[3] == "Critical":
                critVulns.append(elem)
                #print("appended medium element")
            elif  elem[3] == "High":
                highVulns.append(elem)
                #print("appended high element")
            elif elem[3] == "Medium":
                mediumVulns.append(elem)
                #print("appended critical element")
            else:
                pass

    writeOnTemplate(critVulns, documentPar, csvlist)
    writeOnTemplate(highVulns, documentPar, csvlist)
    writeOnTemplate(mediumVulns, documentPar, csvlist)

    mediumVulns.clear()
    critVulns.clear()
    highVulns.clear()

def writeOnTemplate(vulnList, documentPar, csvlist):
    for vuln in vulnList:
        documentPar.insert_paragraph_before("{} - {}".format(vuln[0], vuln[7]), style = vuln[3])#precondition: high, critical, and medium word style must be in the template          
        #if you don't want to use the "high" style you can use the getColor() function
        #documentPar.insert_paragraph_before("{} - {}".format(vuln[0], vuln[7]), style = getColor(vuln[3]))

        #adding text to paragraph
        synopsis = documentPar.insert_paragraph_before()
        synopsis.add_run("Synopsis\n").bold = True
        synopsis.add_run(vuln[8])
        synopsis.add_run("\n")
        
        desc = documentPar.insert_paragraph_before()
        desc.add_run("Description\n").bold = True
        desc.add_run(vuln[9])
        desc.add_run("\n")

        solution = documentPar.insert_paragraph_before()
        solution.add_run("Solution\n").bold = True
        solution.add_run(vuln[10])
        solution.add_run("\n")

        ips = documentPar.insert_paragraph_before()
        ips.add_run("Vulnerable ip(s)\n").bold = True
        for singleIp in getVulnerableIps(copy.deepcopy(csvlist), vuln[0]):
            ips.add_run(singleIp+", ")     
        ips.add_run("\n")


filesPath = sys.argv[1]
filenames = os.listdir(filesPath)


mydoc = docx.Document()
documentTemplate = docx.Document("template.docx")

if filenames:
    for SingleParagraph in documentTemplate.paragraphs:
        if "[VULNS HERE]" in SingleParagraph.text:
            for Singlefile in filenames:
                # print("interacting with " + Singlefile)
                try:
                    file_object = open(filesPath +"/" + Singlefile, "r")
                    csvlist = list(csv.reader(file_object, delimiter = ","))

                    ids = set() #set containing vulns ids (first csv field)

                    writeOnWord(csvlist, ids, SingleParagraph)

                    mydoc.add_page_break()

                    ids.clear() #delete the set (not necessary)
                except Exception as e:
                    print(Singlefile +" doesn't have a csv format compatible with this program\n")
                    # print(e)
                    with open('error.log', 'a') as f:
                        f.write(Singlefile +" doesn't have a csv format compatible with this program\n")
                    pass
            SingleParagraph.clear()
else:
    print("no files in the directory")

documentTemplate.save("result.docx") #save the file